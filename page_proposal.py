"""Lease Proposal (LOI) Generator – choose building, fill info, download .docx."""
import streamlit as st
import os
import io
import datetime
import shared_db

BUILDINGS_INFO = {
    "miami_green": {
        "name": "Miami Green",
        "address": "3150 SW 38th Ave, Miami, FL 33146",
        "management": "Patton Real Estate Management\n3150 SW 38th Ave, Suite 550\nMiami, FL 33146",
        "parking_rate": "$75 per space per month",
    },
    "doral_plaza": {
        "name": "Doral Office Plaza",
        "address": "8550 NW 33rd St, Doral, FL 33122",
        "management": "Patton Real Estate Management\n8550 NW 33rd St, Suite 202\nDoral, FL 33122",
        "parking_rate": "$50 per space per month",
    },
    "waterford": {
        "name": "Waterford Centre",
        "address": "6205 Waterford District, Miami, FL 33126",
        "management": "Patton Real Estate Management\n6205 Waterford District, Suite 310\nMiami, FL 33126",
        "parking_rate": "$85 per space per month",
    },
}


def _get_proposal_history():
    return shared_db.get("proposal_history", [])


def _save_proposal_history(history):
    shared_db.put("proposal_history", history)


def _add_to_proposal_history(tenant, building, suite, fields):
    history = _get_proposal_history()
    entry = {
        "tenant": tenant,
        "building": building,
        "suite": suite,
        "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "fields": dict(fields),
    }
    history = [h for h in history if not (h["tenant"] == tenant and h["building"] == building)]
    history.insert(0, entry)
    _save_proposal_history(history[:50])


def render_proposal_page():
    from utils.style import inject_css, LOGO_B64
    inject_css()

    # Back button
    if st.sidebar.button("Back to Home", key="prop_back"):
        st.session_state["app_mode"] = "home"
        st.rerun()

    # Past Proposals in sidebar
    st.sidebar.markdown(
        '<div style="margin-top:1rem;padding-top:0.5rem;border-top:1px solid rgba(255,255,255,0.08);">'
        '<div style="color:#94a3b8;font-size:0.55rem;letter-spacing:1.5px;text-transform:uppercase;margin-bottom:0.5rem;font-weight:700;">Past Proposals</div>'
        '</div>',
        unsafe_allow_html=True,
    )
    history = _get_proposal_history()
    if history:
        for i, h in enumerate(history):
            label = "{} - {}".format(h["tenant"][:20], h["building"][:15])
            pc1, pc2 = st.sidebar.columns([0.82, 0.18])
            with pc1:
                if st.button(label, key="prop_hist_{}".format(i), use_container_width=True):
                    st.session_state["_load_proposal"] = i
                    st.rerun()
            with pc2:
                if st.button("X", key="del_prop_{}".format(i)):
                    history.pop(i)
                    _save_proposal_history(history)
                    st.rerun()
    else:
        st.sidebar.markdown(
            '<div style="font-size:0.72rem;color:#64748b;font-style:italic;">No proposals yet</div>',
            unsafe_allow_html=True,
        )

    logo_tag = ""
    if LOGO_B64:
        logo_tag = '<img src="data:image/png;base64,{}" style="height:50px;">'.format(LOGO_B64)
    st.markdown(
        '<div style="display:flex;align-items:center;gap:1rem;margin-bottom:1.5rem;">'
        '{logo}'
        '<div><h2 style="margin:0;color:var(--navy);">Lease Proposal Generator</h2>'
        '<div style="font-size:0.75rem;color:var(--slate-400);">Create a professional LOI for any Patton building</div></div>'
        '</div>'.format(logo=logo_tag),
        unsafe_allow_html=True,
    )

    # Check if loading from history
    load_idx = st.session_state.pop("_load_proposal", None)
    if load_idx is not None and load_idx < len(history):
        loaded = history[load_idx]
        f = loaded.get("fields", {})
        st.markdown(
            '<div class="card" style="padding:1.2rem;margin-bottom:1rem;">'
            '<div style="font-size:1.1rem;font-weight:700;color:var(--text-primary);margin-bottom:0.8rem;">Proposal Summary</div>'
            '<table style="width:100%;border-collapse:collapse;">'
            '<tr><td style="padding:0.4rem 0.8rem;font-weight:600;color:var(--text-tertiary);font-size:0.8rem;width:35%;">Tenant</td>'
            '<td style="padding:0.4rem 0.8rem;color:var(--text-primary);font-size:0.85rem;">{tenant}</td></tr>'
            '<tr style="background:var(--bg-secondary);"><td style="padding:0.4rem 0.8rem;font-weight:600;color:var(--text-tertiary);font-size:0.8rem;">Building</td>'
            '<td style="padding:0.4rem 0.8rem;color:var(--text-primary);font-size:0.85rem;">{building}</td></tr>'
            '<tr><td style="padding:0.4rem 0.8rem;font-weight:600;color:var(--text-tertiary);font-size:0.8rem;">Suite</td>'
            '<td style="padding:0.4rem 0.8rem;color:var(--text-primary);font-size:0.85rem;">{suite}</td></tr>'
            '<tr style="background:var(--bg-secondary);"><td style="padding:0.4rem 0.8rem;font-weight:600;color:var(--text-tertiary);font-size:0.8rem;">RSF</td>'
            '<td style="padding:0.4rem 0.8rem;color:var(--text-primary);font-size:0.85rem;">{rsf}</td></tr>'
            '<tr><td style="padding:0.4rem 0.8rem;font-weight:600;color:var(--text-tertiary);font-size:0.8rem;">Base Rate</td>'
            '<td style="padding:0.4rem 0.8rem;color:var(--text-primary);font-size:0.85rem;">{rate}</td></tr>'
            '<tr style="background:var(--bg-secondary);"><td style="padding:0.4rem 0.8rem;font-weight:600;color:var(--text-tertiary);font-size:0.8rem;">Term</td>'
            '<td style="padding:0.4rem 0.8rem;color:var(--text-primary);font-size:0.85rem;">{term}</td></tr>'
            '<tr><td style="padding:0.4rem 0.8rem;font-weight:600;color:var(--text-tertiary);font-size:0.8rem;">Date</td>'
            '<td style="padding:0.4rem 0.8rem;color:var(--text-primary);font-size:0.85rem;">{date}</td></tr>'
            '</table></div>'.format(
                tenant=loaded["tenant"], building=loaded["building"],
                suite=loaded["suite"], rsf=f.get("rsf", "N/A"),
                rate=f.get("rate", "N/A"), term=f.get("term", "N/A"),
                date=loaded["date"],
            ),
            unsafe_allow_html=True,
        )
        st.markdown("---")
        st.markdown("Create a new proposal below, or select another from the sidebar.")

    # Step 1: Choose building
    st.markdown("### 1. Select Building")
    cols = st.columns(3)
    building_keys = list(BUILDINGS_INFO.keys())
    sel_key = st.session_state.get("prop_building", None)

    for i, bk in enumerate(building_keys):
        bi = BUILDINGS_INFO[bk]
        with cols[i]:
            selected = sel_key == bk
            border_color = "var(--blue)" if selected else "var(--slate-200)"
            bg = "var(--blue-50)" if selected else "white"
            st.markdown(
                '<div class="card" style="text-align:center;padding:1.5rem 1rem;border:2px solid {bc};background:{bg};cursor:pointer;">'
                '<div style="font-size:1.6rem;margin-bottom:0.5rem;">&#x1F3E2;</div>'
                '<div style="font-size:1rem;font-weight:700;color:var(--navy);">{name}</div>'
                '<div style="font-size:0.72rem;color:var(--slate-500);margin-top:0.3rem;">{addr}</div>'
                '</div>'.format(bc=border_color, bg=bg, name=bi["name"], addr=bi["address"]),
                unsafe_allow_html=True,
            )
            if st.button("Select", key="prop_sel_{}".format(bk), use_container_width=True,
                         type="primary" if selected else "secondary"):
                st.session_state["prop_building"] = bk
                st.rerun()

    if not sel_key:
        st.info("Select a building above to continue.")
        return

    bi = BUILDINGS_INFO[sel_key]

    st.markdown("---")
    st.markdown("### 2. Tenant & Broker Information")

    c1, c2 = st.columns(2)
    with c1:
        tenant_name = st.text_input("Tenant Name (Legal Entity)", key="prop_tenant")
        tenant_broker_name = st.text_input("Tenant Broker Name", key="prop_broker_name")
        tenant_broker_title = st.text_input("Tenant Broker Title", value="Senior Vice President", key="prop_broker_title")
        tenant_broker_company = st.text_input("Tenant Broker Company", key="prop_broker_co")
    with c2:
        tenant_broker_address = st.text_input("Tenant Broker Address", key="prop_broker_addr")
        tenant_broker_phone = st.text_input("Tenant Broker Phone", key="prop_broker_phone")
        tenant_broker_cell = st.text_input("Tenant Broker Cell", key="prop_broker_cell")

    st.markdown("---")
    st.markdown("### 3. Lease Terms")

    c3, c4 = st.columns(2)
    with c3:
        suite_number = st.text_input("Suite Number", key="prop_suite")
        rsf = st.text_input("Rentable Square Feet (RSF)", key="prop_rsf")
        use_desc = st.text_input("Use of Premises", value="general business purposes consistent with the characteristics of a first-class office building", key="prop_use")
        lease_term = st.text_input("Lease Term", value="Five (5) Years", key="prop_term")
        base_rate = st.text_input("Base Rental Rate ($/RSF)", value="$42.00", key="prop_rate")
        annual_increase = st.text_input("Annual Increase", value="four percent (4%)", key="prop_inc")
    with c4:
        base_year = st.text_input("Operating Expenses Base Year", value=str(datetime.date.today().year), key="prop_base_yr")
        ti_allowance = st.text_input("TI Allowance ($/RSF)", value="$50.00", key="prop_ti")
        parking_spaces = st.text_input("Number of Parking Spaces", value="4", key="prop_parking_n")
        security_deposit = st.text_input("Security Deposit", value="To be determined upon financial review", key="prop_deposit")
        proposal_expiration = st.date_input("Proposal Expiration Date",
            value=datetime.date.today() + datetime.timedelta(days=10),
            key="prop_exp")

    st.markdown("---")
    st.markdown("### 4. Landlord Signatories")
    c5, c6 = st.columns(2)
    with c5:
        signatory1_name = st.text_input("Signatory 1 Name", value="William H. Holly", key="prop_sig1")
        signatory1_title = st.text_input("Signatory 1 Title", value="President", key="prop_sig1_t")
    with c6:
        signatory2_name = st.text_input("Signatory 2 Name", value="James McCoy", key="prop_sig2")
        signatory2_title = st.text_input("Signatory 2 Title", value="Associate", key="prop_sig2_t")

    st.markdown("---")

    if st.button("Generate Lease Proposal", type="primary", use_container_width=True, key="prop_generate"):
        if not tenant_name:
            st.error("Please enter the Tenant Name.")
            return
        if not suite_number:
            st.error("Please enter the Suite Number.")
            return

        with st.spinner("Generating Lease Proposal..."):
            docx_bytes = _generate_proposal_docx(
                building=bi,
                tenant_name=tenant_name,
                broker_name=tenant_broker_name,
                broker_title=tenant_broker_title,
                broker_company=tenant_broker_company,
                broker_address=tenant_broker_address,
                broker_phone=tenant_broker_phone,
                broker_cell=tenant_broker_cell,
                suite=suite_number,
                rsf=rsf,
                use=use_desc,
                term=lease_term,
                rate=base_rate,
                increase=annual_increase,
                base_year=base_year,
                ti=ti_allowance,
                parking_spaces=parking_spaces,
                parking_rate=bi["parking_rate"],
                deposit=security_deposit,
                expiration=proposal_expiration.strftime("%B %d, %Y") if proposal_expiration else "",
                sig1_name=signatory1_name,
                sig1_title=signatory1_title,
                sig2_name=signatory2_name,
                sig2_title=signatory2_title,
            )

        filename = "LOI_{}_{}.docx".format(
            tenant_name.replace(" ", "_").upper()[:20],
            bi["name"].replace(" ", "_").upper()[:15],
        )
        # Save to history
        _add_to_proposal_history(
            tenant_name, bi["name"], suite_number,
            {"building_key": sel_key, "tenant": tenant_name, "suite": suite_number,
             "rsf": rsf, "rate": base_rate, "term": lease_term}
        )

        st.download_button(
            label="Download Lease Proposal (.docx)",
            data=docx_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="prop_download",
        )
        st.success("Lease Proposal generated!")

        # Add to Tenant Folder
        st.markdown("---")
        st.markdown("##### Add to Tenant Folder")
        folders = shared_db.get("tenant_folders", {})
        folder_names = list(folders.keys())
        pfc1, pfc2 = st.columns(2)
        with pfc1:
            sel_folder = st.selectbox("Choose folder", [""] + folder_names, key="prop_tf_sel") if folder_names else ""
        with pfc2:
            new_folder = st.text_input("Or create new", value=tenant_name, key="prop_tf_new")
        folder_target = sel_folder if sel_folder else new_folder
        if folder_target and st.button("Save to '{}'".format(folder_target), key="prop_tf_save", type="primary"):
            if folder_target not in folders:
                folders[folder_target] = []
            folders[folder_target].append({
                "type": "Lease Proposal",
                "tenant": tenant_name,
                "filename": filename,
                "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
                "fields": {"building": bi["name"], "suite": suite_number, "rate": base_rate, "term": lease_term},
            })
            shared_db.put("tenant_folders", folders)
            st.success("Saved to tenant folder: {}".format(folder_target))


def _generate_proposal_docx(**kw):
    """Generate LOI .docx matching Patton's format exactly."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(2)

    building = kw["building"]

    # Logo
    logo_path = os.path.join(os.path.dirname(__file__), "image1.png")
    if os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(2.5))

    # Date
    today_str = datetime.date.today().strftime("%B %d, %Y")
    p = doc.add_paragraph(today_str)
    p.paragraph_format.space_before = Pt(12)

    # Broker address block
    if kw.get("broker_name"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(kw["broker_name"])
        run.bold = True
        run.font.size = Pt(10)
        if kw.get("broker_title"):
            p.add_run("\n{}".format(kw["broker_title"])).font.size = Pt(10)
        parts = []
        if kw.get("broker_company"):
            parts.append(kw["broker_company"])
        if kw.get("broker_address"):
            parts.append(kw["broker_address"])
        if kw.get("broker_phone"):
            parts.append("T {}".format(kw["broker_phone"]))
        if kw.get("broker_cell"):
            parts.append("C {}".format(kw["broker_cell"]))
        if parts:
            p.add_run("\n" + " | ".join(parts)).font.size = Pt(9)

    # RE line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("RE: ")
    run.bold = True
    run.font.size = Pt(10)
    p.add_run(
        'Lease Proposal to {} ("Tenant") at {}, {} ("Building" or "Property") '
        'by Patton Real Estate, LLC. ("Landlord")'.format(
            kw["tenant_name"],
            building["name"],
            building["address"],
        )
    ).font.size = Pt(10)

    # Dear line
    if kw.get("broker_name"):
        first = kw["broker_name"].split()[0] if kw["broker_name"] else "Sir/Madam"
        doc.add_paragraph("Dear {}:".format(first))
    else:
        doc.add_paragraph("Dear Sir/Madam:")

    # Intro paragraph
    doc.add_paragraph(
        "Pursuant to our ongoing discussions, we are pleased to present the following terms and "
        "conditions under which the Landlord would consider leasing space to your client."
    )

    # Lease terms as label/value pairs
    def add_term(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("{}:".format(label))
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        # value on new line indented
        p2 = doc.add_paragraph(value)
        p2.paragraph_format.left_indent = Inches(0.5)
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(4)

    add_term("TENANT", kw["tenant_name"])
    add_term("LANDLORD", "Patton Real Estate, LLC")
    add_term("BUILDING", "{}\n{}".format(building["name"], building["address"]))
    add_term("MANAGEMENT", building["management"])
    add_term("USE",
        "Tenant's use of the space shall be for {} in the {} office market.".format(
            kw.get("use", "general business purposes"),
            building["name"]
        ))
    add_term("PREMISES",
        "Suite {} comprising {} Rentable Square Feet of office space.".format(
            kw["suite"], kw.get("rsf", "TBD")))
    add_term("LEASE TERM",
        "The term of the Lease for the Premises shall be {}. "
        "The Lease Commencement Date shall be upon substantial completion of the Premises.".format(
            kw.get("term", "Five (5) Years")))
    add_term("BASE RENTAL RATE",
        "The Base Rental Rate for the Premises for the first 12 months of the Lease Term shall be equal to "
        "{} per rentable square foot, plus applicable sales tax.\n\n"
        "Thereafter, the Base Rental Rate for the Premises shall increase by {} per annum "
        "beginning in the 13th month of the Lease Term.".format(
            kw.get("rate", "$42.00"),
            kw.get("increase", "four percent (4%)")))
    add_term("OPERATING EXPENSES\n& REAL ESTATE TAXES",
        'Tenant shall pay its pro-rata share of Operating Expenses to the extent such expenses '
        'exceed those incurred by Landlord in {} ("Premises Base Year").'.format(
            kw.get("base_year", str(datetime.date.today().year))))
    add_term("SPACE PLANNING /\nTENANT IMPROVEMENTS",
        "The Landlord will permit, bid, and construct the improvements to the Premises with materials "
        "and finishes pursuant to a mutually agreeable space plan and construction documents. "
        "The Landlord shall provide a Tenant Improvement Allowance of up to {} per rentable square foot "
        "for the Premises.".format(kw.get("ti", "$50.00")))
    add_term("PARKING",
        "Tenant will lease in the Building's parking structure {} parking spaces at the Building's "
        "current market rate. Current monthly parking rates are {} plus applicable taxes.".format(
            kw.get("parking_spaces", "4"),
            kw.get("parking_rate", "$75 per space per month")))
    add_term("SIGNAGE",
        "The Landlord shall provide the Tenant Building standard lobby directory signage and, "
        "at Tenant's expense, standard suite entry signage.")
    add_term("SECURITY DEPOSIT", kw.get("deposit", "To be determined upon financial review"))
    add_term("BROKERAGE",
        "Patton Real Estate, LLC is serving as the exclusive leasing agent for {} and will represent "
        "the Landlord.{} This proposal is contingent upon the mutual understanding of Tenant and Landlord "
        "that the Tenant has not dealt with any other brokerage firm in connection with this Proposal{}.".format(
            building["name"],
            " {} will represent the Tenant in this transaction.".format(kw.get("broker_company", "")) if kw.get("broker_company") else "",
            " other than Patton Real Estate, LLC{}".format(
                " and {}".format(kw["broker_company"]) if kw.get("broker_company") else ""
            ),
        ))
    add_term("NON-BINDING",
        "This Proposal constitutes a summary of the basic business terms of a proposed lease agreement "
        "and does not create any contractual obligation by either party to the other. Nothing in this "
        "Proposal shall be construed as creating any rights in favor of either party. This Proposal does "
        "not address all matters upon which agreement must be reached in order for a lease to be consummated.")
    add_term("PROPOSAL EXPIRATION",
        "This Proposal will expire as of {} at 5:00 p.m.".format(
            kw.get("expiration", "TBD")))

    # Closing
    doc.add_paragraph()
    doc.add_paragraph(
        "Thank you in advance for your consideration of this Lease Proposal. "
        "Please do not hesitate to contact us with any questions you may have."
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.add_run("Sincerely,").font.size = Pt(10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("PATTON REAL ESTATE, LLC")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    # Signatories side by side
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    # Remove borders
    from docx.oxml.ns import qn
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn("w:tblPr"), {})
    borders_elem = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders_elem.makeelement(qn("w:{}".format(edge)), {
            qn("w:val"): "none", qn("w:sz"): "0", qn("w:space"): "0", qn("w:color"): "auto"
        })
        borders_elem.append(el)
    tblPr.append(borders_elem)

    # Row 1: names
    c0 = table.cell(0, 0).paragraphs[0]
    r = c0.add_run(kw.get("sig1_name", "William H. Holly"))
    r.bold = True
    r.font.size = Pt(10)
    c1 = table.cell(0, 1).paragraphs[0]
    r = c1.add_run(kw.get("sig2_name", "James McCoy"))
    r.bold = True
    r.font.size = Pt(10)
    # Row 2: titles
    table.cell(1, 0).paragraphs[0].add_run(kw.get("sig1_title", "President")).font.size = Pt(10)
    table.cell(1, 1).paragraphs[0].add_run(kw.get("sig2_title", "Associate")).font.size = Pt(10)

    # Acceptance section
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("AGREED TO AND ACCEPTED BY:")
    run.bold = True
    run.font.size = Pt(10)

    for label in ["NAME: {}".format(kw["tenant_name"]), "TITLE:", "DATE:"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.add_run("_________________________________________").font.size = Pt(10)
        p2 = doc.add_paragraph(label)
        p2.paragraph_format.space_before = Pt(2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
