"""Lease Abstract Generator – upload scanned PDF, auto-extract, generate .docx."""
import streamlit as st
import os
import io
import datetime
import threading
import shared_db


def _get_scan_history():
    """Get scan history from shared DB."""
    return shared_db.get("scan_history", [])


def _save_scan_history(history):
    """Save scan history to shared DB."""
    shared_db.put("scan_history", history)


def _add_to_scan_history(fields, filename, page_images=None):
    """Save a scan result to history."""
    history = _get_scan_history()
    tenant = fields.get("tenant_name", "").strip() or "Unknown Tenant"
    prop = fields.get("property_name", "").strip() or "Unknown Property"
    suite = fields.get("suite", "").strip()
    entry = {
        "tenant": tenant,
        "property": prop,
        "suite": suite,
        "filename": filename,
        "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "fields": dict(fields),
    }
    # Store page images in session state only (too large for JSON)
    if page_images:
        st.session_state["_last_page_images"] = page_images
    # Avoid duplicates by filename
    history = [h for h in history if h["filename"] != filename]
    history.insert(0, entry)
    _save_scan_history(history[:50])


def render_abstract_page():
    from utils.style import inject_css, LOGO_B64
    inject_css()

    # ── Sidebar ──
    if st.sidebar.button("Back to Home", key="abs_back"):
        st.session_state["app_mode"] = "home"
        st.rerun()

    # Past Scans in sidebar
    st.sidebar.markdown(
        '<div style="margin-top:1rem;padding-top:0.5rem;border-top:1px solid rgba(255,255,255,0.08);">'
        '<div style="color:#94a3b8;font-size:0.55rem;letter-spacing:1.5px;text-transform:uppercase;margin-bottom:0.5rem;font-weight:700;">Past Lease Scans</div>'
        '</div>',
        unsafe_allow_html=True,
    )
    history = _get_scan_history()
    if history:
        for i, h in enumerate(history):
            label = "{} - {}".format(h["tenant"], h.get("suite", ""))
            sc1, sc2 = st.sidebar.columns([0.82, 0.18])
            with sc1:
                if st.button(label, key="scan_hist_{}".format(i), use_container_width=True):
                    st.session_state["_load_scan"] = i
                    st.rerun()
            with sc2:
                if st.button("X", key="del_scan_{}".format(i)):
                    history.pop(i)
                    _save_scan_history(history)
                    st.rerun()
    else:
        st.sidebar.markdown(
            '<div style="font-size:0.72rem;color:#64748b;font-style:italic;">No scans yet</div>',
            unsafe_allow_html=True,
        )

    logo_tag = ""
    if LOGO_B64:
        logo_tag = '<img src="data:image/png;base64,{}" style="height:50px;">'.format(LOGO_B64)
    st.markdown(
        '<div style="display:flex;align-items:center;gap:1rem;margin-bottom:1.5rem;">'
        '{logo}'
        '<div><h2 style="margin:0;color:var(--text-primary);">Lease Abstract Generator</h2>'
        '<div style="font-size:0.75rem;color:var(--text-muted);">Upload a scanned lease PDF &rarr; auto-analysis &rarr; download abstract</div></div>'
        '</div>'.format(logo=logo_tag),
        unsafe_allow_html=True,
    )

    # Check if loading from history
    load_idx = st.session_state.pop("_load_scan", None)
    if load_idx is not None and load_idx < len(history):
        loaded = history[load_idx]
        st.markdown(
            '<div class="card" style="padding:1.2rem;margin-bottom:1rem;">'
            '<div style="font-size:1rem;font-weight:700;color:var(--text-primary);">{tenant}</div>'
            '<div style="font-size:0.78rem;color:var(--text-tertiary);margin-top:0.3rem;">'
            '{prop} | Suite {suite} | {date}</div>'
            '<div style="font-size:0.72rem;color:var(--text-muted);margin-top:0.2rem;">File: {fn}</div>'
            '</div>'.format(
                tenant=loaded["tenant"], prop=loaded["property"],
                suite=loaded.get("suite", "N/A"), date=loaded["date"],
                fn=loaded["filename"],
            ),
            unsafe_allow_html=True,
        )
        # Show scanned PDF pages if available
        page_images = loaded.get("page_images", [])
        if page_images:
            with st.expander("Scanned PDF Pages ({} pages)".format(len(page_images)), expanded=True):
                for pi, img in enumerate(page_images):
                    st.image(img, caption="Page {}".format(pi + 1), use_container_width=True)
        fields = loaded["fields"]
        _render_edit_fields(fields, loaded["filename"])
        return

    # ── Background scan status ─────────────────────────
    bg = st.session_state.get("_bg_scan")
    if bg and bg.get("done"):
        st.success("Scan complete: **{}** — click below to view results.".format(bg.get("filename", "")))
        if st.button("View Scan Results", type="primary", key="abs_view_bg"):
            st.session_state["_bg_scan_show"] = True
            st.rerun()
    elif bg and not bg.get("done"):
        st.info("Scanning **{}** in background... You can use other sections and come back.".format(bg.get("filename", "")))
        if st.button("Refresh Status", key="abs_refresh_bg"):
            st.rerun()

    # Show background scan results
    if st.session_state.pop("_bg_scan_show", False) and bg and bg.get("done"):
        pdf_text = bg.get("text", "")
        page_images = bg.get("images", [])
        filename = bg.get("filename", "scan.pdf")
        st.session_state["_bg_scan"] = None  # Clear
        has_text = pdf_text and len(pdf_text.strip()) > 50
        if has_text:
            st.success("Text extracted ({} chars). Auto-parsing lease terms...".format(len(pdf_text)))
            with st.expander("Raw Extracted Text", expanded=False):
                st.text_area("", pdf_text, height=300, key="abs_raw_text_bg")
            fields = _parse_lease_fields(pdf_text)
            try:
                _add_to_scan_history(fields, filename, page_images)
            except Exception:
                pass
        else:
            st.warning("No text extracted. Fill in fields manually.")
            if page_images:
                with st.expander("Page Previews ({} pages)".format(len(page_images)), expanded=True):
                    for i, img in enumerate(page_images):
                        st.image(img, caption="Page {}".format(i + 1), use_container_width=True)
            fields = {}
        _render_edit_fields(fields, filename)
        return

    uploaded = st.file_uploader(
        "Drop your scanned lease PDF here",
        type=["pdf"],
        key="abs_pdf_upload",
    )

    if uploaded is None:
        st.markdown(
            '<div class="card" style="text-align:center;padding:4rem 2rem;border:2px dashed var(--border);">'
            '<div style="font-size:3rem;margin-bottom:1rem;">&#128196;</div>'
            '<div style="font-size:1.1rem;font-weight:600;color:var(--text-primary);margin-bottom:0.5rem;">Upload a Lease PDF</div>'
            '<div style="font-size:0.82rem;color:var(--text-tertiary);">Drag and drop or click to browse.<br>Supports scanned and digital PDFs.</div>'
            '</div>',
            unsafe_allow_html=True,
        )
        return

    # ── Extract: try fast text first, background OCR if needed ──
    pdf_bytes = uploaded.read()
    uploaded.seek(0)
    filename = uploaded.name

    # Quick text extraction (no OCR)
    pdf_text, page_images = _extract_pdf_text_only(pdf_bytes)
    has_text = pdf_text and len(pdf_text.strip()) > 50

    if has_text:
        st.success("Text extracted ({} chars). Auto-parsing lease terms...".format(len(pdf_text)))
        with st.expander("Raw Extracted Text", expanded=False):
            st.text_area("", pdf_text, height=300, key="abs_raw_text")
        fields = _parse_lease_fields(pdf_text)
        try:
            _add_to_scan_history(fields, filename, page_images)
        except Exception:
            pass
        _render_edit_fields(fields, filename)
    else:
        # Need OCR — launch in background
        st.session_state["_bg_scan"] = {"filename": filename, "done": False, "text": "", "images": []}
        _bg_data = st.session_state["_bg_scan"]
        _bg_bytes = pdf_bytes

        def _run_bg_ocr(data, raw_bytes):
            try:
                text, imgs = _extract_pdf_with_ocr(raw_bytes)
                data["text"] = text
                data["images"] = imgs
            except Exception as e:
                data["error"] = str(e)
            data["done"] = True

        t = threading.Thread(target=_run_bg_ocr, args=(_bg_data, _bg_bytes), daemon=True)
        t.start()
        st.info("No text layer detected. **OCR started in background** — you can navigate to other sections. Come back to see results.")
        if page_images:
            with st.expander("Page Previews", expanded=False):
                for i, img in enumerate(page_images):
                    st.image(img, caption="Page {}".format(i + 1), use_container_width=True)


def _render_edit_fields(fields, filename=""):
    """Render editable lease fields and generate button."""
    has_text = any(v for v in fields.values() if v)

    st.markdown("### Review & Edit Lease Information")
    st.markdown(
        '<div style="font-size:0.78rem;color:var(--text-muted);margin-bottom:1rem;">'
        '{}</div>'.format(
            "Fields auto-filled from PDF. Review, correct if needed, then generate."
            if has_text else
            "Fill in fields manually using the page previews above, then generate."
        ),
        unsafe_allow_html=True,
    )

    # Use a prefix based on whether loaded from history to avoid widget key conflicts
    pfx = "af" if not st.session_state.get("_load_scan_done") else "lf"

    def _v(key, default=""):
        return fields.get(key, default)

    st.markdown("##### 1. Property & Tenant")
    c1, c2 = st.columns(2)
    with c1:
        fields["property_name"] = st.text_input("Property Name", value=_v("property_name"), key=pfx + "_pn")
        fields["property_address"] = st.text_input("Property Address", value=_v("property_address"), key=pfx + "_pa")
        fields["suite"] = st.text_input("Suite No.", value=_v("suite"), key=pfx + "_su")
    with c2:
        fields["tenant_name"] = st.text_input("Tenant Name (Legal Entity)", value=_v("tenant_name"), key=pfx + "_tn")
        fields["trade_name"] = st.text_input("Trade Name (if different)", value=_v("trade_name"), key=pfx + "_td")
        fields["industry"] = st.text_input("Industry / Use", value=_v("industry"), key=pfx + "_in")

    st.markdown("##### 2. Lease Overview")
    c3, c4 = st.columns(2)
    with c3:
        fields["lease_type"] = st.text_input("Lease Type", value=_v("lease_type"), key=pfx + "_lt")
        fields["lease_date"] = st.text_input("Lease Date", value=_v("lease_date"), key=pfx + "_ld")
        fields["lease_term"] = st.text_input("Lease Term", value=_v("lease_term"), key=pfx + "_lm")
        fields["commencement"] = st.text_input("Commencement Date", value=_v("commencement"), key=pfx + "_cd")
    with c4:
        fields["rent_commencement"] = st.text_input("Rent Commencement", value=_v("rent_commencement"), key=pfx + "_rc")
        fields["expiration"] = st.text_input("Expiration Date", value=_v("expiration"), key=pfx + "_ex")
        fields["option_period"] = st.text_input("Option / Renewal", value=_v("option_period"), key=pfx + "_op")

    st.markdown("##### 3. Financial Terms")
    fields["rent_schedule"] = st.text_area(
        "Rent Schedule (one line per year: Period | Yearly | Monthly)",
        value=_v("rent_schedule"),
        height=150, key=pfx + "_rs",
        help="e.g. Year 1 (Jan 2025 - Dec 2025) | $120,000.00 | $10,000.00",
    )
    c5, c6 = st.columns(2)
    with c5:
        fields["operating_expenses"] = st.text_input("Operating Expenses / CAM", value=_v("operating_expenses"), key=pfx + "_oe")
        fields["base_year_cam"] = st.text_input("Base Year - CAM", value=_v("base_year_cam"), key=pfx + "_by")
        fields["additional_rent"] = st.text_input("Additional Rent", value=_v("additional_rent"), key=pfx + "_ar")
    with c6:
        fields["annual_increase"] = st.text_input("Annual Percentage Increase", value=_v("annual_increase"), key=pfx + "_ai")
        fields["security_deposit"] = st.text_input("Security Deposit", value=_v("security_deposit"), key=pfx + "_sd")
        fields["late_fees"] = st.text_input("Late Fees / Interest", value=_v("late_fees"), key=pfx + "_lf")

    st.markdown("##### 4. Key Dates & Critical Deadlines")
    c7, c8 = st.columns(2)
    with c7:
        fields["notice_dates"] = st.text_input("Notice Dates (Renewal/Termination)", value=_v("notice_dates"), key=pfx + "_nd")
        fields["escalation_dates"] = st.text_input("Rent Escalation Dates", value=_v("escalation_dates"), key=pfx + "_ed")
    with c8:
        fields["free_rent"] = st.text_input("Free Rent Period", value=_v("free_rent"), key=pfx + "_fr")

    st.markdown("##### 5. TI & Build Out")
    c9, c10 = st.columns(2)
    with c9:
        fields["ti_allowance"] = st.text_area("TI Allowance", value=_v("ti_allowance"), height=80, key=pfx + "_ti")
    with c10:
        fields["scope_responsibility"] = st.text_area("Scope Responsibility", value=_v("scope_responsibility"), height=80, key=pfx + "_sr")

    st.markdown("##### 6. Maintenance & Repairs")
    fields["landlord_resp"] = st.text_area("Landlord Responsibilities", value=_v("landlord_resp"), height=80, key=pfx + "_lr")
    fields["tenant_resp"] = st.text_area("Tenant Responsibilities", value=_v("tenant_resp"), height=80, key=pfx + "_tr")
    c11, c12 = st.columns(2)
    with c11:
        fields["hvac_resp"] = st.text_input("HVAC Responsibility", value=_v("hvac_resp"), key=pfx + "_hv")
    with c12:
        fields["after_hours_hvac"] = st.text_input("After Hours HVAC Charges", value=_v("after_hours_hvac"), key=pfx + "_ah")

    st.markdown("##### 7. Parking & Access")
    c13, c14 = st.columns(2)
    with c13:
        fields["parking_spaces"] = st.text_input("Number of Spaces", value=_v("parking_spaces"), key=pfx + "_ps")
        fields["parking_reserved"] = st.text_input("Reserved / Unreserved", value=_v("parking_reserved"), key=pfx + "_pr")
    with c14:
        fields["parking_rate"] = st.text_input("Monthly Rate", value=_v("parking_rate"), key=pfx + "_pm")

    # ── Generate & Download ─────────────────────────────
    st.markdown("---")
    tenant_name = fields.get("tenant_name", "").strip() or "Unknown Tenant"
    fname = "LEASE_ABSTRACT_{}.docx".format(
        tenant_name.replace(" ", "_").upper()[:30]
    )
    try:
        docx_bytes = _generate_abstract_docx(fields)
        st.download_button(
            "Generate & Download Lease Abstract (.docx)", docx_bytes, fname,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True, key=pfx + "_dl", type="primary",
        )
    except Exception as e:
        st.error("Error generating abstract: {}".format(e))
        docx_bytes = None

    # Add to Tenant Folder option
    st.markdown("---")
    st.markdown("##### Add to Tenant Folder")

    folders = shared_db.get("tenant_folders", {})
    folder_names = list(folders.keys())

    fc1, fc2 = st.columns(2)
    with fc1:
        if folder_names:
            selected_folder = st.selectbox("Choose existing folder", [""] + folder_names, key=pfx + "_tf_sel")
        else:
            selected_folder = ""
            st.markdown('<div style="font-size:0.78rem;color:var(--text-muted);">No folders yet</div>', unsafe_allow_html=True)
    with fc2:
        new_folder = st.text_input("Or create new folder", value=tenant_name, key=pfx + "_tf_new")

    folder_target = selected_folder if selected_folder else new_folder
    if folder_target and st.button("Save to '{}'".format(folder_target), key=pfx + "_tf_save", type="primary"):
        if folder_target not in folders:
            folders[folder_target] = []
        doc_entry = {
            "type": "Lease Abstract",
            "tenant": tenant_name,
            "filename": fname,
            "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
            "fields": dict(fields),
        }
        # Store docx bytes in session for download from tenant folder
        if docx_bytes:
            if "_tenant_docx" not in st.session_state:
                st.session_state["_tenant_docx"] = {}
            st.session_state["_tenant_docx"][fname] = docx_bytes
        folders[folder_target].append(doc_entry)
        shared_db.put("tenant_folders", folders)
        st.success("Saved to tenant folder: {}".format(folder_target))


# ═══════════════════════════════════════════════════════
# PDF extraction
# ═══════════════════════════════════════════════════════

def _extract_pdf_text_only(pdf_bytes):
    """Fast text extraction only (no OCR). Returns (text, page_images)."""
    try:
        import fitz
    except ImportError:
        return "", []
    text_parts = []
    images = []
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for i, page in enumerate(doc):
            text_parts.append(page.get_text())
            if i < 5:
                mat = fitz.Matrix(1.5, 1.5)
                pix = page.get_pixmap(matrix=mat)
                images.append(pix.tobytes("png"))
        doc.close()
    except Exception:
        return "", []
    return "\n".join(text_parts), images


def _extract_pdf_with_ocr(pdf_bytes):
    """Full extraction with OCR fallback (runs in background thread)."""
    try:
        import fitz
    except ImportError:
        return "", []
    text_parts = []
    images = []
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for i, page in enumerate(doc):
        text_parts.append(page.get_text())
        if i < 5:
            mat = fitz.Matrix(1.5, 1.5)
            pix = page.get_pixmap(matrix=mat)
            images.append(pix.tobytes("png"))

    combined = "\n".join(text_parts)
    if len(combined.strip()) < 50:
        try:
            import pytesseract
            from PIL import Image as PILImage
            total_pages = len(doc)
            page_indices = list(range(min(5, total_pages)))
            if total_pages > 7:
                page_indices += list(range(total_pages - 2, total_pages))
            seen = set()
            page_indices = [p for p in page_indices if not (p in seen or seen.add(p))]
            ocr_parts = []
            for i in page_indices:
                page = doc[i]
                pix = page.get_pixmap(dpi=150)
                img = PILImage.frombytes("RGB", [pix.width, pix.height], pix.samples)
                page_text = pytesseract.image_to_string(img)
                ocr_parts.append(page_text)
            combined = "\n\n".join(ocr_parts)
        except ImportError:
            pass
        except Exception:
            pass
    doc.close()
    return combined, images


# ═══════════════════════════════════════════════════════
# Auto-parse lease text
# ═══════════════════════════════════════════════════════

def _parse_lease_fields(text):
    """Parse lease fields from extracted text (works with OCR or digital text)."""
    import re
    f = {}
    t = text

    def _find(patterns, default=""):
        """Try multiple regex patterns, return first match."""
        if isinstance(patterns, str):
            patterns = [patterns]
        for pattern in patterns:
            m = re.search(pattern, t, re.IGNORECASE | re.DOTALL)
            if m:
                val = m.group(1).strip()
                val = re.sub(r'\s+', ' ', val)
                # Clean OCR artifacts
                val = val.strip('.,;: ')
                return val[:500]
        return default

    # Property & Tenant
    f["property_name"] = _find([
        r"(?:property\s*name|building\s*name)[:\s]*(.+?)(?:\n|$)",
        r"(?:building\s*(?:located|known)\s*(?:at|as))\s+(.+?)(?:\n|$)",
    ])
    f["property_address"] = _find([
        r"(?:property\s*address|building\s*address)[:\s]*(.+?)(?:\n|$)",
        r"(?:premises|building)\s*(?:located|known)\s*(?:at|as)\s+(.+?)(?:\n|$)",
        r"(?:located\s+at)\s+(\d+[^,\n]+(?:,\s*\w+)+)",
    ])
    f["suite"] = _find([
        r"[Ss]uite\s*(?:[Nn]o\.?\s*)?(\w+\d+\w*)",
        r"[Pp]remises[:\s]*[Ss]uite\s*(?:[Nn]o\.?\s*)?(\w+)",
    ])
    f["tenant_name"] = _find([
        r"[Tt]enant[:\s]*([A-Z][A-Z\s&,.]+(?:LLC|INC|CORP|LP|LLP|CO|COMPANY|SOLUTIONS|GROUP|SERVICES)[.,\s]*(?:a\s+Florida[^,\n]*)?)",
        r"[Tt]enant[:\s]*([^\n,]+(?:LLC|Inc|Corp))",
        r"(?:and|&)\s+([A-Z][A-Z\s]+(?:LLC|INC|CORP))\s*[,(].*?[Tt]enant",
    ])
    f["trade_name"] = _find(r"(?:trade\s*name|d/?b/?a)[:\s]*(.+?)(?:\n|$)")
    f["industry"] = _find([
        r"[Uu]se\s*(?:of\s*)?[Pp]remises[:\s]*(.+?)(?:\n|$)",
        r"[Pp]urpose[:\s]*(.+?)(?:\n|$)",
        r"(?:use|purpose)[:\s]+([A-Z][a-z].*?)(?:\.|$)",
    ])

    # Lease Overview
    f["lease_type"] = _find([
        r"(?:lease\s*type|type\s*of\s*lease)[:\s]*(.+?)(?:\n|$)",
        r"(OFFICE\s*LEASE|COMMERCIAL\s*LEASE|RETAIL\s*LEASE)",
    ])
    f["lease_date"] = _find([
        r"[Dd]ate\s*(?:of\s*)?[Ll]ease[:\s]*(.+?)(?:\n|$)",
        r"[Dd]ated[:\s]*(.+?)(?:\n|$)",
        r"(?:made|entered)\s*(?:as\s*of|on)\s*(?:the\s*)?(.+?)(?:\s*by|\n|$)",
    ])
    f["lease_term"] = _find([
        r"[Tt]erm[:\s]*([^\n]*?(?:year|month|yr|mo)[s]?[^\n]*)",
        r"(?:initial\s*)?[Tt]erm[:\s]*(.+?)(?:\n|$)",
    ])
    f["commencement"] = _find([
        r"[Ll]ease\s*[Cc]ommencement\s*[Dd]ate[:\s]*(.+?)(?:\n|$)",
        r"[Cc]ommencement\s*[Dd]ate[:\s]*(.+?)(?:\n|$)",
        r"[Ss]tart\s*[Dd]ate[:\s]*(.+?)(?:\n|$)",
    ])
    f["rent_commencement"] = _find([
        r"[Rr]ent\s*[Cc]ommencement\s*[Dd]ate[:\s]*(.+?)(?:\n|$)",
    ])
    f["expiration"] = _find([
        r"[Ee]xpiration\s*[Dd]ate[:\s]*(.+?)(?:\n|$)",
        r"[Tt]ermination\s*[Dd]ate[:\s]*(.+?)(?:\n|$)",
    ])
    f["option_period"] = _find([
        r"[Oo]ption\s*(?:to\s*)?[Rr]enew[:\s]*(.+?)(?:\n|$)",
        r"[Rr]enewal[:\s]*(.+?)(?:\n|$)",
    ])

    # Financial
    f["rent_schedule"] = _extract_rent_schedule(t)
    f["operating_expenses"] = _find([
        r"[Oo]perating\s*[Ee]xpense[s]?[:\s]*(.+?)(?:\n|$)",
        r"CAM[:\s]*(.+?)(?:\n|$)",
        r"[Cc]ommon\s*[Aa]rea\s*[Mm]aintenance[:\s]*(.+?)(?:\n|$)",
    ])
    f["base_year_cam"] = _find(r"[Bb]ase\s*[Yy]ear[:\s]*(.+?)(?:\n|$)")
    f["additional_rent"] = _find([
        r"[Aa]dditional\s*[Rr]ent[:\s]*(.+?)(?:\n|$)",
        r"[Aa]dditional\s*[Rr]ent[:\s]*(.+?(?:\$[\d,.]+).+?)(?:\n\n|\n[A-Z])",
    ])
    f["annual_increase"] = _find([
        r"(?:annual\s*(?:percentage\s*)?increase|escalation)[:\s]*(.+?)(?:\n|$)",
        r"(\d+(?:\.\d+)?%?\s*(?:per\s*annum|annual|yearly))",
    ])
    f["security_deposit"] = _find([
        r"[Ss]ecurity\s*[Dd]eposit[:\s]*(.+?)(?:\n|$)",
        r"[Ss]ecurity\s*[Dd]eposit[:\s]*(\$[\d,.]+)",
    ])
    f["late_fees"] = _find([
        r"[Ll]ate\s*[Ff]ee[s]?[:\s]*(.+?)(?:\n|$)",
        r"[Ll]ate\s*[Cc]harge[s]?[:\s]*(.+?)(?:\n|$)",
    ])

    # Key Dates
    f["notice_dates"] = _find([
        r"[Nn]otice\s*(?:period|date|required)[:\s]*(.+?)(?:\n|$)",
        r"(\d+)\s*(?:days?|months?)\s*(?:prior\s*)?(?:written\s*)?notice",
    ])
    f["escalation_dates"] = _find(r"[Ee]scalation\s*[Dd]ate[:\s]*(.+?)(?:\n|$)")
    f["free_rent"] = _find([
        r"[Ff]ree\s*[Rr]ent[:\s]*(.+?)(?:\n|$)",
        r"[Rr]ent\s*[Aa]batement[:\s]*(.+?)(?:\n|$)",
    ])

    # TI
    f["ti_allowance"] = _find([
        r"[Tt]enant\s*[Ii]mprovement\s*[Aa]llowance[:\s]*(.+?)(?:\n\n|\n[A-Z])",
        r"TI\s*[Aa]llowance[:\s]*(.+?)(?:\n|$)",
        r"[Ii]mprovement\s*[Aa]llowance[:\s]*(\$[\d,.]+[^.\n]*)",
    ])
    f["scope_responsibility"] = _find(r"(?:scope|build.?out)\s*responsibilit[yies]*[:\s]*(.+?)(?:\n\n|\n[A-Z])")

    # Maintenance
    f["landlord_resp"] = _find(r"[Ll]andlord\s*(?:shall|responsibilit)[:\s]*(.+?)(?:\n\n|\n[A-Z])")
    f["tenant_resp"] = _find(r"[Tt]enant\s*(?:shall|responsibilit)[:\s]*(.+?)(?:\n\n|\n[A-Z])")
    f["hvac_resp"] = _find(r"HVAC[:\s]*(.+?)(?:\n|$)")
    f["after_hours_hvac"] = _find(r"[Aa]fter\s*[Hh]ours?\s*HVAC[:\s]*(.+?)(?:\n|$)")

    # Parking
    f["parking_spaces"] = _find([
        r"(?:[Nn]umber\s*of\s*)?(?:[Pp]arking\s*)?[Ss]paces[:\s]*(\d+)",
        r"[Pp]arking[:\s]*(\d+\s*(?:spaces|stalls))",
        r"[Dd]esignated\s*[Pp]arking\s*[Ss]paces[:\s]*(\d+)",
        r"[Uu]nassigned\s*[Ss]paces[:\s]*(\d+)",
    ])
    f["parking_reserved"] = _find([
        r"((?:\d+\s*)?[Rr]eserved[^.\n]*)",
        r"((?:\d+\s*)?[Uu]nreserved[^.\n]*)",
        r"[Dd]esignated[:\s]*(\d+).*[Uu]nassigned[:\s]*(\d+)",
    ])
    f["parking_rate"] = _find(r"[Pp]arking\s*(?:monthly\s*)?[Rr]ate[:\s]*(.+?)(?:\n|$)")

    # Rentable area (bonus field for context)
    rsf = _find([
        r"[Rr]entable\s*[Aa]rea[:\s]*([0-9,]+\s*(?:rentable\s*)?(?:square\s*feet|RSF|sf))",
        r"([0-9,]+)\s*rentable\s*square\s*feet",
    ])
    if rsf and not f.get("property_address"):
        f["property_address"] = rsf

    return f


def _extract_rent_schedule(text):
    """Try to find a rent schedule table in the text."""
    import re
    lines = text.split("\n")
    schedule_lines = []
    in_schedule = False
    for line in lines:
        line = line.strip()
        if not line:
            if in_schedule:
                break
            continue
        # Look for lines with dollar amounts that look like rent rows
        has_dollar = bool(re.search(r'\$[\d,]+\.?\d*', line))
        has_period = bool(re.search(r'(?:yr\s*\d|year\s*\d|\d{4}\s*[-–]|month\s*\d|period\s*\d|[1-9]\s*[-–]\s*\d)', line, re.IGNORECASE))
        if has_dollar and has_period:
            in_schedule = True
            schedule_lines.append(line)
        elif in_schedule and has_dollar:
            schedule_lines.append(line)
        elif in_schedule and re.match(r'^\d', line):
            schedule_lines.append(line)
        elif in_schedule:
            break
    return "\n".join(schedule_lines)


# ═══════════════════════════════════════════════════════
# Generate .docx matching Patton format exactly
# ═══════════════════════════════════════════════════════

def _generate_abstract_docx(fields):
    """Generate lease abstract .docx matching the Patton table format."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    NAVY = RGBColor(0x0F, 0x17, 0x2A)
    GRAY = RGBColor(0x47, 0x55, 0x69)
    BLUE = RGBColor(0x25, 0x63, 0xEB)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    LABEL_BG = "F1F5F9"
    HEADER_BG = "1E3A5F"

    # ── Logo ──
    logo_path = os.path.join(os.path.dirname(__file__), "abstract_header.jpg")
    if not os.path.exists(logo_path):
        logo_path = os.path.join(os.path.dirname(__file__), "image1.png")
    if os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Inches(3.5))
        p.paragraph_format.space_after = Pt(6)

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LEASE ABSTRACT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    title.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("{} | {}".format(
        fields.get("tenant_name", "Tenant") or "Tenant",
        fields.get("property_name", "Property") or "Property",
    ))
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY
    sub.paragraph_format.space_after = Pt(10)

    def _shade(cell, color):
        shd = cell._element.get_or_add_tcPr().makeelement(qn("w:shd"), {
            qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color,
        })
        cell._element.get_or_add_tcPr().append(shd)

    def _set_cell_margins(cell, top=40, bot=40, left=80, right=80):
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        mar = tcPr.makeelement(qn("w:tcMar"), {})
        for side, val in [("top", top), ("bottom", bot), ("left", left), ("right", right)]:
            el = mar.makeelement(qn("w:{}".format(side)), {qn("w:w"): str(val), qn("w:type"): "dxa"})
            mar.append(el)
        tcPr.append(mar)

    def add_section_table(num, title_text, rows_data):
        """Create a table with header row spanning 2 cols + label|value rows."""
        n_rows = 1 + len(rows_data)
        tbl = doc.add_table(rows=n_rows, cols=2)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        for row in tbl.rows:
            row.cells[0].width = Inches(2.8)
            row.cells[1].width = Inches(4.6)

        # Header row - merge and style
        hdr = tbl.cell(0, 0)
        hdr_r = tbl.cell(0, 1)
        hdr.merge(hdr_r)
        p = hdr.paragraphs[0]
        run = p.add_run("{}. {}".format(num, title_text))
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = WHITE
        _shade(hdr, HEADER_BG)
        _set_cell_margins(hdr, 60, 60, 100, 100)

        # Data rows
        for i, (label, value) in enumerate(rows_data):
            ri = i + 1
            cl = tbl.cell(ri, 0)
            cr = tbl.cell(ri, 1)
            # Label
            pl = cl.paragraphs[0]
            rl = pl.add_run(label)
            rl.bold = True
            rl.font.size = Pt(9)
            rl.font.color.rgb = GRAY
            _shade(cl, LABEL_BG)
            _set_cell_margins(cl, 30, 30, 80, 80)
            # Value
            pr = cr.paragraphs[0]
            rv = pr.add_run(str(value) if value else "N/A")
            rv.font.size = Pt(9.5)
            rv.font.color.rgb = NAVY
            _set_cell_margins(cr, 30, 30, 80, 80)

        # Spacer
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(6)
        sp.paragraph_format.space_after = Pt(2)

    def add_rent_schedule_table(schedule_text):
        """Add the 3-column rent schedule table if data exists."""
        if not schedule_text or not schedule_text.strip():
            return
        lines = [l.strip() for l in schedule_text.strip().split("\n") if l.strip()]
        if not lines:
            return

        # Section header
        tbl_hdr = doc.add_table(rows=1, cols=1)
        tbl_hdr.style = "Table Grid"
        hdr = tbl_hdr.cell(0, 0)
        p = hdr.paragraphs[0]
        run = p.add_run("Base Rent Schedule (Year-by-Year / Step Rent)")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        _shade(hdr, LABEL_BG)

        # Parse lines into rows
        import re
        parsed = []
        for line in lines:
            parts = re.split(r'\s*\|\s*', line)
            if len(parts) >= 3:
                parsed.append(parts[:3])
            elif len(parts) == 2:
                parsed.append([parts[0], parts[1], ""])
            else:
                # Try splitting by tab or multiple spaces
                parts = re.split(r'\t+|\s{2,}', line)
                if len(parts) >= 3:
                    parsed.append(parts[:3])
                elif len(parts) == 2:
                    parsed.append([parts[0], parts[1], ""])
                else:
                    parsed.append([line, "", ""])

        if not parsed:
            return

        n = len(parsed)
        tbl = doc.add_table(rows=n + 1, cols=3)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for ci, txt in enumerate(["Period", "Yearly Payment", "Monthly Payment"]):
            cell = tbl.cell(0, ci)
            p = cell.paragraphs[0]
            run = p.add_run(txt)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE
            _shade(cell, HEADER_BG)
            _set_cell_margins(cell, 30, 30, 60, 60)

        for ri, row_data in enumerate(parsed):
            for ci in range(3):
                val = row_data[ci] if ci < len(row_data) else ""
                cell = tbl.cell(ri + 1, ci)
                p = cell.paragraphs[0]
                run = p.add_run(val.strip())
                run.font.size = Pt(9)
                run.font.color.rgb = NAVY
                _set_cell_margins(cell, 25, 25, 60, 60)

        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(4)

    # ══ Build sections ══

    # 1. Property & Tenant
    add_section_table(1, "Property & Tenant Information", [
        ("Property Name", fields.get("property_name")),
        ("Property Address", fields.get("property_address")),
        ("Suite No.", fields.get("suite")),
        ("Tenant Name (Legal Entity)", fields.get("tenant_name")),
        ("Trade Name (if different)", fields.get("trade_name")),
        ("Industry / Use", fields.get("industry")),
    ])

    # 2. Lease Overview
    add_section_table(2, "Lease Overview", [
        ("Lease Type", fields.get("lease_type")),
        ("Lease Date", fields.get("lease_date")),
        ("Lease Term", fields.get("lease_term")),
        ("Commencement Date", fields.get("commencement")),
        ("Rent Commencement Date", fields.get("rent_commencement")),
        ("Lease Expiration Date", fields.get("expiration")),
        ("Option Period", fields.get("option_period")),
    ])

    # 3. Financial Terms
    # Header-only table for section 3
    tbl3h = doc.add_table(rows=1, cols=1)
    tbl3h.style = "Table Grid"
    hdr3 = tbl3h.cell(0, 0)
    p3 = hdr3.paragraphs[0]
    r3 = p3.add_run("3. Financial Terms")
    r3.bold = True
    r3.font.size = Pt(11)
    r3.font.color.rgb = WHITE
    _shade(hdr3, HEADER_BG)
    _set_cell_margins(hdr3, 60, 60, 100, 100)

    # Rent schedule table
    add_rent_schedule_table(fields.get("rent_schedule", ""))

    # Financial detail rows
    fin_rows = [
        ("Operating Expenses (CAM)", fields.get("operating_expenses")),
        ("Base Year - CAM", fields.get("base_year_cam")),
        ("Additional Rent", fields.get("additional_rent")),
        ("Annual Percentage Increase", fields.get("annual_increase")),
        ("Security Deposit", fields.get("security_deposit")),
        ("Late Fees / Interest", fields.get("late_fees")),
    ]
    fin_rows = [(k, v) for k, v in fin_rows if v]
    if fin_rows:
        tbl_fin = doc.add_table(rows=len(fin_rows), cols=2)
        tbl_fin.style = "Table Grid"
        tbl_fin.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (label, value) in enumerate(fin_rows):
            cl = tbl_fin.cell(i, 0)
            cr = tbl_fin.cell(i, 1)
            cl.width = Inches(2.8)
            cr.width = Inches(4.6)
            pl = cl.paragraphs[0]
            rl = pl.add_run(label)
            rl.bold = True
            rl.font.size = Pt(9)
            rl.font.color.rgb = GRAY
            _shade(cl, LABEL_BG)
            _set_cell_margins(cl, 30, 30, 80, 80)
            pr = cr.paragraphs[0]
            rv = pr.add_run(str(value) if value else "N/A")
            rv.font.size = Pt(9.5)
            rv.font.color.rgb = NAVY
            _set_cell_margins(cr, 30, 30, 80, 80)

    doc.add_paragraph().paragraph_format.space_before = Pt(6)

    # 4. Key Dates
    dates_rows = [
        ("Notice Dates (Renewal / Termination)", fields.get("notice_dates")),
        ("Rent Escalation Dates", fields.get("escalation_dates")),
        ("Free Rent Period", fields.get("free_rent")),
    ]
    dates_rows = [(k, v) for k, v in dates_rows if v]
    if dates_rows:
        add_section_table(4, "Key Dates & Critical Deadlines", dates_rows)

    # 5. TI & Build Out
    ti_rows = [
        ("TI Allowance", fields.get("ti_allowance")),
        ("Scope Responsibility", fields.get("scope_responsibility")),
    ]
    ti_rows = [(k, v) for k, v in ti_rows if v]
    if ti_rows:
        add_section_table(5, "Tenant Improvement (TI) & Build Out", ti_rows)

    # 6. Maintenance
    maint_rows = [
        ("Landlord Responsibilities", fields.get("landlord_resp")),
        ("Tenant Responsibilities", fields.get("tenant_resp")),
        ("HVAC Responsibility", fields.get("hvac_resp")),
        ("After Hours HVAC Charges", fields.get("after_hours_hvac")),
    ]
    maint_rows = [(k, v) for k, v in maint_rows if v]
    if maint_rows:
        add_section_table(6, "Maintenance & Repairs", maint_rows)

    # 7. Parking
    park_rows = [
        ("Number of Spaces", fields.get("parking_spaces")),
        ("Reserved / Unreserved", fields.get("parking_reserved")),
        ("Monthly Rate", fields.get("parking_rate")),
    ]
    park_rows = [(k, v) for k, v in park_rows if v]
    if park_rows:
        add_section_table(7, "Parking & Access", park_rows)

    # Footer
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Prepared by Patton Real Estate Mgmt LLC | {}".format(
        datetime.date.today().strftime("%B %d, %Y")
    ))
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
